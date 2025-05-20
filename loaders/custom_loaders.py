"""
커스텀 문서 로더
다양한 파일 형식을 지원하는 최적화된 로더
"""
import json
import pandas as pd
import markdown
from pathlib import Path
from typing import List
import logging

import docx
from llama_index.core import Document, SimpleDirectoryReader

logger = logging.getLogger(__name__)

class CustomDocumentLoader:
    """다양한 문서 형식을 지원하는 커스텀 로더"""
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.csv', '.md', '.docx', '.json'}
    
    def load_documents(self, file_paths: List[Path]) -> List[Document]:
        """
        다중 파일 로드
        
        Args:
            file_paths: 로드할 파일 경로 리스트
            
        Returns:
            Document 객체 리스트
        """
        documents = []
        
        for file_path in file_paths:
            try:
                if file_path.suffix not in self.supported_extensions:
                    logger.warning(f"지원하지 않는 파일 형식: {file_path}")
                    continue
                
                # 파일별 로드 메소드 호출
                if file_path.suffix == '.pdf':
                    docs = self._load_pdf(file_path)
                elif file_path.suffix == '.csv':
                    docs = self._load_csv(file_path)
                elif file_path.suffix == '.md':
                    docs = self._load_markdown(file_path)
                elif file_path.suffix == '.docx':
                    docs = self._load_docx(file_path)
                elif file_path.suffix == '.json':
                    docs = self._load_json(file_path)
                else:
                    continue
                
                documents.extend(docs)
                logger.debug(f"로드 완료: {file_path} ({len(docs)}개 청크)")
                
            except Exception as e:
                logger.error(f"파일 로드 실패 {file_path}: {e}")
                continue
        
        logger.info(f"총 {len(documents)}개 문서 청크 로드 완료")
        return documents
    
    def _load_pdf(self, file_path: Path) -> List[Document]:
        """PDF 파일 로드"""
        try:
            reader = SimpleDirectoryReader(input_files=[str(file_path)])
            documents = reader.load_data()
            
            # 메타데이터 추가
            for doc in documents:
                doc.metadata.update({
                    'source': file_path.name,
                    'file_type': 'pdf',
                    'file_path': str(file_path)
                })
            
            return documents
        except Exception as e:
            logger.error(f"PDF 로드 실패 {file_path}: {e}")
            return []
    
    def _load_csv(self, file_path: Path) -> List[Document]:
        """CSV 파일 로드 (특정 컬럼 선택 가능)"""
        try:
            df = pd.read_csv(file_path)
            
            # CSV가 비어있는지 확인
            if df.empty:
                logger.warning(f"빈 CSV 파일: {file_path}")
                return []
            
            # 선택할 컬럼 결정 (Name, Purchase 우선, 없으면 모든 컬럼)
            preferred_columns = ["Name", "Purchase"]
            selected_columns = [col for col in preferred_columns if col in df.columns]
            
            if not selected_columns:
                selected_columns = df.columns.tolist()
                logger.info(f"기본 컬럼 사용: {selected_columns}")
            
            # 데이터를 텍스트로 변환
            csv_text = df[selected_columns].to_string(index=False)
            
            document = Document(
                text=csv_text,
                metadata={
                    'source': file_path.name,
                    'file_type': 'csv',
                    'file_path': str(file_path),
                    'columns': selected_columns,
                    'row_count': len(df)
                }
            )
            
            return [document]
            
        except Exception as e:
            logger.error(f"CSV 로드 실패 {file_path}: {e}")
            return []
    
    def _load_markdown(self, file_path: Path) -> List[Document]:
        """Markdown 파일 로드"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Markdown을 HTML로 변환
            html_content = markdown.markdown(md_content)
            
            document = Document(
                text=html_content,
                metadata={
                    'source': file_path.name,
                    'file_type': 'markdown',
                    'file_path': str(file_path),
                    'original_format': 'markdown'
                }
            )
            
            return [document]
            
        except Exception as e:
            logger.error(f"Markdown 로드 실패 {file_path}: {e}")
            return []
    
    def _load_docx(self, file_path: Path) -> List[Document]:
        """DOCX 파일 로드"""
        try:
            doc = docx.Document(file_path)
            
            # 모든 문단 텍스트 추출
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            docx_text = "\n".join(paragraphs)
            
            if not docx_text.strip():
                logger.warning(f"빈 DOCX 파일: {file_path}")
                return []
            
            document = Document(
                text=docx_text,
                metadata={
                    'source': file_path.name,
                    'file_type': 'docx',
                    'file_path': str(file_path),
                    'paragraph_count': len(paragraphs)
                }
            )
            
            return [document]
            
        except Exception as e:
            logger.error(f"DOCX 로드 실패 {file_path}: {e}")
            return []
    
    def _load_json(self, file_path: Path) -> List[Document]:
        """JSON 파일 로드"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
            
            # JSON을 보기 좋은 텍스트로 변환
            json_text = json.dumps(json_data, ensure_ascii=False, indent=2)
            
            document = Document(
                text=json_text,
                metadata={
                    'source': file_path.name,
                    'file_type': 'json',
                    'file_path': str(file_path),
                    'json_size': len(json_text)
                }
            )
            
            return [document]
            
        except Exception as e:
            logger.error(f"JSON 로드 실패 {file_path}: {e}")
            return []